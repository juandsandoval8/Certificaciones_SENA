import tkinter as tk
from tkinter import filedialog, messagebox, ttk
from tkcalendar import DateEntry
import pandas as pd
from docx import Document
import os
from datetime import datetime

df = None
excel_path = None
plantilla_path = None

def replace_tags(doc, replacements):
    for p in doc.paragraphs:
        inline_text = p.text
        for key, value in replacements.items():
            if key in inline_text:
                inline_text = inline_text.replace(key, value)
        for run in p.runs:
            run.text = ""
        if p.runs:
            p.runs[0].text = inline_text
        else:
            p.add_run(inline_text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    inline_text = p.text
                    for key, value in replacements.items():
                        if key in inline_text:
                            inline_text = inline_text.replace(key, value)
                    for run in p.runs:
                        run.text = ""
                    if p.runs:
                        p.runs[0].text = inline_text
                    else:
                        p.add_run(inline_text)

def cargar_excel():
    global df, excel_path
    excel_path = filedialog.askopenfilename(title="Selecciona el archivo Excel", filetypes=[("Excel files", "*.xlsx")])
    if excel_path:
        df = pd.read_excel(excel_path)
        # Limpieza de nombres de columnas para evitar problemas con espacios
        df.columns = df.columns.str.strip()
        # Convertir columnas de fecha a datetime si existen
        for col in ['fecha de contrato', 'Fecha de inicio', 'Fecha fin']:
            if col in df.columns:
                df[col] = pd.to_datetime(df[col], errors='coerce')
        combo_nombre['values'] = df['Nombre'].tolist()
        messagebox.showinfo("Éxito", "Excel cargado correctamente.")

def cargar_plantilla():
    global plantilla_path
    plantilla_path = filedialog.askopenfilename(title="Selecciona la plantilla Word", filetypes=[("Word files", "*.docx")])
    if plantilla_path:
        messagebox.showinfo("Éxito", "Plantilla cargada correctamente.")

def autocompletar(event):
    nombre_sel = combo_nombre.get()
    datos = df[df['Nombre'] == nombre_sel].iloc[0]

    entry_cedula.delete(0, tk.END)
    entry_cedula.insert(0, str(datos['cedula']))

    entry_lugar_expedicion.delete(0, tk.END)
    entry_lugar_expedicion.insert(0, datos['lugar de expedicion'])

    entry_num_contrato.delete(0, tk.END)
    entry_num_contrato.insert(0, str(datos['Numero de contrato']))

    if pd.notnull(datos['fecha de contrato']):
        fecha_contrato.set_date(datos['fecha de contrato'])
    else:
        fecha_contrato.set_date(datetime.today())

    entry_objeto.delete(0, tk.END)
    entry_objeto.insert(0, datos['Objeto'])

    entry_plazo_ejecucion.delete(0, tk.END)
    entry_plazo_ejecucion.insert(0, str(datos['Plazo de ejecución']))

    if pd.notnull(datos['Fecha de inicio']):
        fecha_inicio.set_date(datos['Fecha de inicio'])
    else:
        fecha_inicio.set_date(datetime.today())

    if pd.notnull(datos['Fecha fin']):
        fecha_fin.set_date(datos['Fecha fin'])
    else:
        fecha_fin.set_date(datetime.today())

    entry_termino_ejecucion.delete(0, tk.END)
    entry_termino_ejecucion.insert(0, str(datos['Termino de ejecución']))

    entry_valor.delete(0, tk.END)
    entry_valor.insert(0, str(datos['Valor']))

    entry_obligaciones.delete(0, tk.END)
    entry_obligaciones.insert(0, datos['Obligaciones'])

def generar_doc():
    if not plantilla_path or not excel_path:
        messagebox.showerror("Error", "Carga primero el Excel y la plantilla.")
        return

    nombre = combo_nombre.get()
    cedula = entry_cedula.get()
    lugar_expedicion = entry_lugar_expedicion.get()
    num_contrato = entry_num_contrato.get()
    fecha_contrato_val = fecha_contrato.get_date().strftime("%d/%m/%Y")
    objeto = entry_objeto.get()
    plazo_ejecucion = entry_plazo_ejecucion.get()
    fecha_inicio_val = fecha_inicio.get_date().strftime("%d/%m/%Y")
    fecha_fin_val = fecha_fin.get_date().strftime("%d/%m/%Y")
    termino_ejecucion = entry_termino_ejecucion.get()
    valor = entry_valor.get()
    obligaciones = entry_obligaciones.get()
    fecha_expedicion_val = datetime.today().strftime("%d/%m/%Y")

    replacements = {
        "[Nombre]": nombre,
        "[cedula]": cedula,
        "[lugar de expedición]": lugar_expedicion,
        "[número de contrato]": num_contrato,
        "[fecha del contrato]": fecha_contrato_val,
        "[Objeto]": objeto,
        "[número de días]": plazo_ejecucion,
        "[fecha de terminación de contrato en formato dd/mm/aaaa ]": fecha_fin_val,
        "[Fecha inicio]": fecha_inicio_val,  # <-- Aquí está el agregado
        "[Termino de ejecución]": termino_ejecucion,
        "[valor en pesos colombianos $, y representación numérica]": valor,
        "[Obligaciones]": obligaciones,
        "[fecha de expedición del mismo día que se genera el documento en formato dd/mm/aaaa ]": fecha_expedicion_val
    }

    doc = Document(plantilla_path)
    replace_tags(doc, replacements)

    output_path = os.path.join(os.path.dirname(plantilla_path), f"Certificado_{nombre}.docx")
    doc.save(output_path)
    messagebox.showinfo("Éxito", f"Documento generado: {output_path}")

root = tk.Tk()
root.title("Generador de Certificados")
root.geometry("750x650")

tk.Button(root, text="Cargar Excel", command=cargar_excel).grid(row=0, column=0, pady=5)
tk.Button(root, text="Cargar Plantilla", command=cargar_plantilla).grid(row=0, column=1, pady=5)

tk.Label(root, text="Nombre:").grid(row=1, column=0, sticky="e")
combo_nombre = ttk.Combobox(root, width=50)
combo_nombre.grid(row=1, column=1)
combo_nombre.bind("<<ComboboxSelected>>", autocompletar)

tk.Label(root, text="Cédula:").grid(row=2, column=0, sticky="e")
entry_cedula = tk.Entry(root, width=53)
entry_cedula.grid(row=2, column=1)

tk.Label(root, text="Lugar de Expedición:").grid(row=3, column=0, sticky="e")
entry_lugar_expedicion = tk.Entry(root, width=53)
entry_lugar_expedicion.grid(row=3, column=1)

tk.Label(root, text="Número de Contrato:").grid(row=4, column=0, sticky="e")
entry_num_contrato = tk.Entry(root, width=53)
entry_num_contrato.grid(row=4, column=1)

tk.Label(root, text="Fecha de Contrato:").grid(row=5, column=0, sticky="e")
fecha_contrato = DateEntry(root, width=50, date_pattern="dd/mm/yyyy")
fecha_contrato.grid(row=5, column=1)

tk.Label(root, text="Objeto:").grid(row=6, column=0, sticky="e")
entry_objeto = tk.Entry(root, width=53)
entry_objeto.grid(row=6, column=1)

tk.Label(root, text="Plazo de Ejecución (días):").grid(row=7, column=0, sticky="e")
entry_plazo_ejecucion = tk.Entry(root, width=53)
entry_plazo_ejecucion.grid(row=7, column=1)

tk.Label(root, text="Fecha de Inicio:").grid(row=8, column=0, sticky="e")
fecha_inicio = DateEntry(root, width=50, date_pattern="dd/mm/yyyy")
fecha_inicio.grid(row=8, column=1)

tk.Label(root, text="Fecha Fin:").grid(row=9, column=0, sticky="e")
fecha_fin = DateEntry(root, width=50, date_pattern="dd/mm/yyyy")
fecha_fin.grid(row=9, column=1)

tk.Label(root, text="Término de Ejecución:").grid(row=10, column=0, sticky="e")
entry_termino_ejecucion = tk.Entry(root, width=53)
entry_termino_ejecucion.grid(row=10, column=1)

tk.Label(root, text="Valor:").grid(row=11, column=0, sticky="e")
entry_valor = tk.Entry(root, width=53)
entry_valor.grid(row=11, column=1)

tk.Label(root, text="Obligaciones:").grid(row=12, column=0, sticky="e")
entry_obligaciones = tk.Entry(root, width=53)
entry_obligaciones.grid(row=12, column=1)

tk.Button(root, text="Generar Documento", command=generar_doc).grid(row=13, column=0, columnspan=2, pady=20)

root.mainloop()
